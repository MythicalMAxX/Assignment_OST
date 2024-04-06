import os
import re
import subprocess
import pdfplumber
import textract
import pandas as pd
from bs4 import BeautifulSoup
from docx import Document
from docx2pdf import convert

def read_text_from_pdf(file_path):
    with pdfplumber.open(file_path) as pdf:
        text = ''
        for page in pdf.pages:
            text += page.extract_text()
    return text

def read_text_from_docx(file_path):
    doc = Document(file_path)
    text = ''
    for paragraph in doc.paragraphs:
        text += paragraph.text + '\n'
    return text

def convert_to_pdf(input_file_path, output_file_path):
    subprocess.run(['libreoffice', '--headless', '--convert-to', 'pdf', input_file_path, '--outdir', os.path.dirname(output_file_path)])

def convert_doc_to_docx(input_file_path, output_file_path):
    with open(input_file_path, 'rb') as f:
        html_content = f.read()
        soup = BeautifulSoup(html_content, 'lxml')
        text = soup.get_text()
    doc = Document()
    doc.add_paragraph(text)
    doc.save(output_file_path)

def extract_email(text):
    text = text.decode('utf-8') if isinstance(text, bytes) else text
    email = re.findall(r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b', text)
    return email[0] if email else None

def extract_phone(text):
    text = text.decode('utf-8') if isinstance(text, bytes) else text
    phone = re.findall(r'\+?\d{1,4}?\s?\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}', text)
    return phone[0] if phone else None

def clean_text(text):
    text = text.decode('utf-8') if isinstance(text, bytes) else text
    text = re.sub('\s+', ' ', text)
    text = re.sub('\n+', '\n', text)
    return text.strip()

def process_files(folder_path):
    data = []
    for filename in os.listdir(folder_path):
        file_path = os.path.join(folder_path, filename)
        file_name, file_extension = os.path.splitext(filename)
        if file_extension.lower() == '.pdf':
            text = read_text_from_pdf(file_path)
        elif file_extension.lower() == '.docx':
            text = read_text_from_docx(file_path)
        elif file_extension.lower() == '.doc':
            docx_path = os.path.join(folder_path, f"{file_name}.docx")
            convert_doc_to_docx(file_path, docx_path)
            convert(docx_path, os.path.join(folder_path, f"{file_name}.pdf"))
            text = read_text_from_pdf(os.path.join(folder_path, f"{file_name}.pdf"))
            if len(text) < 10:
                text = textract.process(os.path.join(folder_path, f"{file_name}.doc"))
                print(text)
            os.remove(docx_path)
            os.remove(os.path.join(folder_path, f"{file_name}.pdf"))
        else:
            print(f"Unsupported file type: {filename}")
            continue
        email = extract_email(text)
        phone = extract_phone(text)
        text = clean_text(text)
        data.append([filename, email, phone, text])
        print(f"Text from {filename}:")
        print(text)
    return data

def write_to_excel(data, output_file):
    df = pd.DataFrame(data, columns=['Filename', 'Email', 'Phone', 'Text'])
    df.to_excel(output_file, index=False)

folder_path = "CV/Sample2/"
output_file = "output.xlsx"
data = process_files(folder_path)
write_to_excel(data, output_file)
